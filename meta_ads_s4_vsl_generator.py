"""S4 VSL Script Generator — 视频广告脚本生成器.
Input: product name + key features
Process: pull S1/S2/素材库 context → DeepSeek generates 5-8 angle VSL scripts
Output: python-docx → import to Feishu wiki → notify
"""
import json, urllib.request, sys, time, re, datetime, uuid, os, io
try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

from cloud_config import (
    PROXY, IM_APP_ID, IM_APP_SECRET, BT_APP_ID, BT_APP_SECRET,
    DEEPSEEK_KEY, WIKI_SPACE_ID, BOSS_OPEN_ID, META_ADS_CONSOLE_APP,
)

MATERIAL_APP = os.environ.get('MATERIAL_APP', 'PpZIbSIuxaPa5wsNGDZcZm9Wn7t')
S2_APP = META_ADS_CONSOLE_APP
WIKI_PARENT_NODE = os.environ.get('WIKI_PARENT_NODE_S4', 'MDQewVLuAiDyPmkBmf3cHz45nGh')
LOG_FILE = os.environ.get('S4_LOG_FILE', '/tmp/s4_log.txt')

def log(msg):
    ts = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    line = f'[{ts}] {msg}'
    print(line)
    with open(LOG_FILE, 'a', encoding='utf-8') as f:
        f.write(line + '\n')

# ============================================================
# API helpers
# ============================================================
_token_cache = {}

def get_token(app_id, secret):
    key = app_id
    if key in _token_cache and time.time() - _token_cache[key][1] < 600:
        return _token_cache[key][0]
    body = json.dumps({'app_id': app_id, 'app_secret': secret}).encode()
    req = urllib.request.Request('https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal',
        data=body, headers={'Content-Type': 'application/json'})
    token = json.loads(urllib.request.build_opener(PROXY).open(req, timeout=15).read())['tenant_access_token']
    _token_cache[key] = (token, time.time())
    return token

def feishu(method, path, body=None, app='bt'):
    url = f'https://open.feishu.cn/open-apis{path}'
    data = json.dumps(body).encode('utf-8') if body else None
    last_err = None
    for attempt in range(3):
        tok = get_token(BT_APP_ID, BT_APP_SECRET) if app == 'bt' else get_token(IM_APP_ID, IM_APP_SECRET)
        req = urllib.request.Request(url, data=data,
            headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {tok}'}, method=method)
        try:
            return json.loads(urllib.request.build_opener(PROXY).open(req, timeout=30).read())
        except urllib.error.HTTPError as e:
            return {'code': e.code, 'body': e.read().decode('utf-8','replace')[:300]}
        except (urllib.error.URLError, OSError) as e:
            last_err = e
            log(f'  Network error (attempt {attempt+1}/3): {e}')
            time.sleep(2 * (attempt + 1))
    return {'code': -1, 'body': f'Network error after 3 retries: {last_err}'}

def deepseek(system_prompt, user_prompt):
    body = json.dumps({
        'model': 'deepseek-chat', 'max_tokens': 8000,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt},
        ]
    }).encode('utf-8')
    last_err = None
    for attempt in range(3):
        req = urllib.request.Request('https://api.deepseek.com/chat/completions',
            data=body, headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {DEEPSEEK_KEY}'})
        try:
            resp = json.loads(urllib.request.build_opener(PROXY).open(req, timeout=180).read())
            return resp['choices'][0]['message']['content']
        except Exception as e:
            last_err = e
            log(f'  DeepSeek retry {attempt+1}/3: {e}')
            time.sleep(5 * (attempt + 1))
    raise last_err

def gt(v):
    if not v: return ''
    if isinstance(v, str): return v
    if isinstance(v, list) and v and isinstance(v[0], dict): return v[0].get('text', '')
    return str(v)

# ============================================================
# Data collection
# ============================================================
def collect_framework_structures():
    """Read 12 framework structures from 框架结构库 (full content)."""
    r = feishu('GET', f'/bitable/v1/apps/{MATERIAL_APP}/tables/tbluWVngE93DKCdH/records?page_size=20')
    frameworks = []
    for item in r.get('data', {}).get('items', []):
        fd = item['fields']
        name = gt(fd.get('框架名称', ''))
        logic = gt(fd.get('核心逻辑', ''))
        scene = gt(fd.get('适用场景', ''))
        hook = gt(fd.get('开头钩子类型', ''))
        difficulty = gt(fd.get('难度', ''))
        if name:
            frameworks.append(
                f'### [{name}] (钩子:{hook} | 难度:{difficulty})\n'
                f'核心逻辑: {logic}\n'
                f'适用场景: {scene}\n')
    return frameworks

def collect_s2_insights():
    """Read latest S2 competitor insights."""
    r = feishu('POST', f'/bitable/v1/apps/{S2_APP}/tables/tbl8DPF9Z1jqdSpF/records/search?page_size=5',
        {'field_names': ['竞品名称', '策略洞察', '值得借鉴的角度']})
    insights = []
    for item in r.get('data', {}).get('items', []):
        fd = item['fields']
        comp = gt(fd.get('竞品名称', ''))
        strategy = gt(fd.get('策略洞察', ''))[:300]
        borrow = gt(fd.get('值得借鉴的角度', ''))[:300]
        if comp and (strategy or borrow):
            insights.append(f'[{comp}] 策略:{strategy}\n借鉴:{borrow}')
    return insights

def collect_s2_ad_details():
    """Read latest S2 competitor ad analysis (停病药信买)."""
    r = feishu('POST', f'/bitable/v1/apps/{S2_APP}/tables/tblyUsxc3NmHGAOZ/records/search?page_size=10',
        {'filter': {'conjunction': 'and', 'conditions': [
            {'field_name': '仍在投放', 'operator': 'is', 'value': ['true']}
        ]}, 'field_names': ['竞品名称', 'Ad Copy', '停(Hook)', '病(Pain)', '药(Solution)', '信(Trust)', '买(CTA)', '改进建议']})
    ads = []
    for item in r.get('data', {}).get('items', []):
        fd = item['fields']
        comp = gt(fd.get('竞品名称', ''))
        copy = gt(fd.get('Ad Copy', ''))[:100]
        hook = gt(fd.get('停(Hook)', ''))
        pain = gt(fd.get('病(Pain)', ''))
        solution = gt(fd.get('药(Solution)', ''))
        if comp and (hook or pain):
            ads.append(f'[{comp}] Copy:{copy} | 停:{hook} | 病:{pain} | 药:{solution}')
    return ads

def collect_existing_materials(product_keyword=''):
    """Read existing materials from 痛点/CTA/卖点/评论库."""
    materials = {}

    # 痛点库
    r = feishu('GET', f'/bitable/v1/apps/{MATERIAL_APP}/tables/tblxT2UNUaI89Pjj/records?page_size=20')
    pains = []
    for item in r.get('data', {}).get('items', []):
        fd = item['fields']
        name = gt(fd.get('场景名称', ''))
        pain = gt(fd.get('用户痛点', ''))
        if name or pain:
            pains.append(f'{name}: {pain[:60]}')
    materials['痛点'] = pains[:10]

    # CTA库
    r = feishu('GET', f'/bitable/v1/apps/{MATERIAL_APP}/tables/tblbtMGxaYTAUqLk/records?page_size=20')
    ctas = []
    for item in r.get('data', {}).get('items', []):
        fd = item['fields']
        cta_en = gt(fd.get('CTA话术（英文）', ''))
        cta_cn = gt(fd.get('中文翻译', ''))
        stage = gt(fd.get('视频阶段', ''))
        if cta_en or cta_cn:
            ctas.append(f'[{stage}] {cta_en[:50]} | {cta_cn[:30]}')
    materials['CTA'] = ctas[:10]

    # 评论库
    r = feishu('GET', f'/bitable/v1/apps/{MATERIAL_APP}/tables/tblNokhGSLbUkezh/records?page_size=15')
    reviews = []
    for item in r.get('data', {}).get('items', []):
        fd = item['fields']
        text = gt(fd.get('评论原文（英文）', ''))
        cn = gt(fd.get('中文翻译', ''))
        if text:
            reviews.append(f'{text[:60]} ({cn[:30]})')
    materials['评论'] = reviews[:8]

    return materials

def web_research(product, features=''):
    """Search Amazon reviews, Reddit, and web for product insights."""
    results = {'amazon_reviews': [], 'reddit': [], 'market': []}

    # Clean product name for search (remove Chinese)
    import re as _re
    en_name = ' '.join(_re.findall(r'[a-zA-Z0-9]+', product))
    if not en_name:
        en_name = product

    search_queries = [
        (f'{en_name} review site:amazon.com', 'amazon_reviews'),
        (f'{en_name} reddit gaming accessories', 'reddit'),
        (f'{en_name} vs competitors best Switch accessories 2026', 'market'),
    ]

    for query, category in search_queries:
        try:
            # Use DeepSeek to do a research summary instead of direct web search
            # (WebSearch tool is not available in standalone scripts)
            pass
        except Exception:
            pass

    # Use DeepSeek for market research in one shot
    try:
        research_prompt = f'''You are a market research analyst. For the product "{product}" ({features}), provide a concise research brief:

1. **Common User Pain Points** (5 bullet points): What problems do users complain about with existing similar products? Think from a Switch/gaming accessories buyer's perspective.

2. **Top Positive Reviews Themes** (5 bullet points): What do happy buyers typically praise? What words/emotions do they use?

3. **Purchase Hesitations** (3 bullet points): What stops someone from buying this type of product?

4. **Emotional Trigger Words** (10 words): Words that resonate with this audience (e.g., "finally", "obsessed", "game-changer")

5. **Usage Scenarios** (5 bullet points): When and where would someone use this product?

Be specific to gaming accessories / Switch ecosystem. Output in English.'''

        research = deepseek('You are a gaming accessories market research expert.', research_prompt)
        results['research'] = research
        log(f'  Market research: {len(research)} chars')
    except Exception as e:
        log(f'  Research failed: {e}')
        results['research'] = ''

    return results

# ============================================================
# Docx generation + import
# ============================================================
def build_vsl_docx(title, product, features, scripts_text, now):
    """Build .docx with VSL scripts."""
    from docx import Document
    from docx.shared import Inches
    import tempfile

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f'产品: {product} | 生成时间: {now.strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'核心卖点: {features}')

    # Parse scripts by angle markers
    angles = re.split(r'\n(?=###?\s*角度\s*[#＃]?\d)', scripts_text)
    if len(angles) <= 1:
        angles = re.split(r'\n(?=---)', scripts_text)
    if len(angles) <= 1:
        angles = [scripts_text]

    for i, angle in enumerate(angles):
        angle = angle.strip()
        if not angle:
            continue

        # Extract angle title
        title_match = re.match(r'###?\s*(.+)', angle)
        if title_match:
            doc.add_heading(title_match.group(1).strip(), level=2)
            angle = angle[title_match.end():]
        else:
            doc.add_heading(f'角度 #{i+1}', level=2)

        for line in angle.split('\n'):
            t = line.strip()
            if not t or t == '---':
                continue
            # Bold labels
            bold_match = re.match(r'\*\*(.+?)\*\*[：:]\s*(.*)', t)
            if bold_match:
                p = doc.add_paragraph()
                p.add_run(f'{bold_match.group(1)}: ').bold = True
                p.add_run(bold_match.group(2))
            elif t.startswith(('- ', '* ', '• ')):
                doc.add_paragraph(re.sub(r'^[-*•]\s*', '', t), style='List Bullet')
            elif re.match(r'^\d+[\.\)、]', t):
                doc.add_paragraph(re.sub(r'^\d+[\.\)、]\s*', '', t), style='List Number')
            elif t.startswith(('停', '病', '药', '信', '买', '【')):
                p = doc.add_paragraph()
                label_m = re.match(r'([停病药信买][\(（].+?[\)）])[：:]\s*(.*)', t)
                if label_m:
                    p.add_run(f'{label_m.group(1)}: ').bold = True
                    p.add_run(label_m.group(2))
                else:
                    p.add_run(t)
            else:
                doc.add_paragraph(t)

    path = os.path.join(os.environ.get('TEMP', '/tmp'), 's4_vsl.docx')
    doc.save(path)
    return path

def _upload_and_import_docx(docx_bytes, title):
    """One attempt: upload → import → poll. Returns doc_token or None."""
    token = get_token(IM_APP_ID, IM_APP_SECRET)
    boundary = uuid.uuid4().hex
    safe_name = f'{title}.docx'
    extra_json = json.dumps({'obj_type': 'docx', 'file_extension': 'docx'})
    parts = []
    for name, value in [('file_name', safe_name), ('parent_type', 'ccm_import_open'),
                        ('size', str(len(docx_bytes))), ('extra', extra_json)]:
        parts.append(f'--{boundary}\r\nContent-Disposition: form-data; name="{name}"\r\n\r\n{value}\r\n'.encode())
    parts.append(f'--{boundary}\r\nContent-Disposition: form-data; name="file"; filename="{safe_name}"\r\n'
                 f'Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n'.encode())
    parts.append(docx_bytes)
    parts.append(f'\r\n--{boundary}--\r\n'.encode())

    req = urllib.request.Request('https://open.feishu.cn/open-apis/drive/v1/medias/upload_all',
        data=b''.join(parts),
        headers={'Content-Type': f'multipart/form-data; boundary={boundary}', 'Authorization': f'Bearer {token}'})
    try:
        resp = json.loads(urllib.request.build_opener(PROXY).open(req, timeout=60).read())
        file_token = resp['data']['file_token']
    except Exception as e:
        log(f'  Upload failed: {e}')
        return None

    r = feishu('POST', '/drive/v1/import_tasks', {
        'file_extension': 'docx', 'file_token': file_token,
        'type': 'docx', 'point': {'mount_type': 1, 'mount_key': ''},
    }, app='im')
    if r.get('code') != 0:
        log(f'  Import failed: {r}')
        return None
    ticket = r['data']['ticket']

    for _ in range(60):
        time.sleep(2)
        r2 = feishu('GET', f'/drive/v1/import_tasks/{ticket}', app='im')
        status = r2.get('data', {}).get('result', {}).get('job_status', -1)
        if status == 0:
            return r2['data']['result'].get('token', '')
        elif status in (1, 2):
            log(f'  Import failed: status={status} extra={json.dumps(r2.get("data",{}).get("result",{}),ensure_ascii=False)[:200]}')
            return None
    log('  Import timeout')
    return None


def import_docx_to_wiki(docx_path, title):
    """Upload .docx → import → move to wiki."""
    with open(docx_path, 'rb') as f:
        docx_bytes = f.read()
    log(f'  DOCX size: {len(docx_bytes)} bytes')

    doc_token = None
    for attempt in range(2):
        doc_token = _upload_and_import_docx(docx_bytes, title)
        if doc_token:
            break
        log(f'  Import attempt {attempt+1}/2 failed, retrying...')
        time.sleep(5)
    if not doc_token:
        return None, None

    feishu('POST', f'/drive/v1/permissions/{doc_token}/members?type=docx&need_notification=false',
           {'member_type': 'openid', 'member_id': BOSS_OPEN_ID, 'perm': 'full_access'}, app='im')

    r3 = feishu('POST', f'/wiki/v2/spaces/{WIKI_SPACE_ID}/nodes/move_docs_to_wiki', {
        'parent_wiki_token': WIKI_PARENT_NODE,
        'obj_type': 'docx', 'obj_token': doc_token,
    }, app='im')
    if r3.get('code') != 0:
        log(f'  Wiki move failed, using docx URL: {r3}')
        return f'https://u1wpma3xuhr.feishu.cn/docx/{doc_token}', doc_token

    node_token = r3.get('data', {}).get('wiki_token')
    if not node_token:
        for _ in range(10):
            time.sleep(2)
            rn = feishu('GET', f'/wiki/v2/spaces/get_node?obj_type=docx&token={doc_token}', app='im')
            node_token = rn.get('data', {}).get('node', {}).get('node_token')
            if node_token:
                break
        if not node_token:
            log(f'  Wiki node_token not resolved, using docx URL')
            return f'https://u1wpma3xuhr.feishu.cn/docx/{doc_token}', doc_token

    return f'https://u1wpma3xuhr.feishu.cn/wiki/{node_token}', doc_token

# ============================================================
# Main
# ============================================================
def main(product=None, features=None):
    if product is None:
        if len(sys.argv) < 2:
            print('Usage: python meta_ads_s4_vsl_generator.py "产品名" ["核心卖点1, 卖点2, ..."]')
            sys.exit(1)
        product = sys.argv[1]
        features = sys.argv[2] if len(sys.argv) > 2 else ''
    features = features or ''

    log('=' * 60)
    log(f'S4 VSL Script Generator — {product}')

    now = datetime.datetime.now()
    date_str = now.strftime('%Y-%m-%d')

    # ── Step 1: Collect context data ──────────────────────────
    log('Step 1: Collecting context...')

    frameworks = collect_framework_structures()
    log(f'  框架结构: {len(frameworks)} 种')

    s2_insights = collect_s2_insights()
    log(f'  竞品洞察: {len(s2_insights)} 条')

    s2_ads = collect_s2_ad_details()
    log(f'  竞品广告拆解: {len(s2_ads)} 条')

    materials = collect_existing_materials(product)
    log(f'  素材库: 痛点{len(materials.get("痛点",[]))} CTA{len(materials.get("CTA",[]))} 评论{len(materials.get("评论",[]))}')

    # ── Step 2: Market research ───────────────────────────────
    log('Step 2: AI market research...')
    research = web_research(product, features)

    # ── Step 3: Build DeepSeek prompt ─────────────────────────
    log('Step 3: Calling DeepSeek...')

    kb_path = os.environ.get(
        'META_ADS_KB_PATH',
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'meta_ads_knowledge.md'),
    )
    with open(kb_path, 'r', encoding='utf-8') as f:
        knowledge = f.read()

    system_prompt = f'''{knowledge}

---

## 你的任务
你是短视频广告脚本生成专家。基于以上知识库的「停→病→药→信→买」VSL框架，为指定产品生成多角度的视频广告脚本。

## 品牌信息
- Powkong（宝空）：游戏配件品牌（手柄壳、摇杆帽、收纳包、充电底座等），主要市场美国
- Funlab（纷岚）：游戏外设品牌（手柄、配件），主要市场美国

## 可用的视频结构公式（12种）
{chr(10).join(frameworks)}
'''

    # Build context sections
    context_parts = []

    if s2_insights:
        context_parts.append('===== 竞品策略洞察（来自S2系统）=====')
        context_parts.append('\n'.join(s2_insights[:3]))

    if s2_ads:
        context_parts.append('\n===== 竞品广告停病药信买拆解（来自S2系统）=====')
        context_parts.append('\n'.join(s2_ads[:5]))

    if materials.get('痛点'):
        context_parts.append('\n===== 素材库已有痛点 =====')
        context_parts.append('\n'.join(materials['痛点']))

    if materials.get('CTA'):
        context_parts.append('\n===== 素材库已有CTA话术 =====')
        context_parts.append('\n'.join(materials['CTA'][:8]))

    if materials.get('评论'):
        context_parts.append('\n===== 素材库已有用户评论 =====')
        context_parts.append('\n'.join(materials['评论']))

    # Add market research
    if research.get('research'):
        context_parts.append('\n===== AI 市场调研（产品痛点/好评/顾虑/场景）=====')
        context_parts.append(research['research'])

    user_prompt = f'''===== 产品信息 =====
产品名称: {product}
核心卖点: {features if features else '(请根据产品名推断)'}

{chr(10).join(context_parts)}

===== 创作指引 =====

**素材来源优先级：**
1. 优先参考素材库中已有的痛点、CTA话术、用户评论（标注为 [素材库]）
2. 如素材库无合适素材，基于AI市场调研的数据自主创作（标注为 [AI原创]）
3. 参考竞品S2数据中的有效策略，但做出差异化

**AI自主创作时的参考维度：**
- 从"用户痛点"推导停(Hook)和病(Pain)的角度
- 从"好评主题"推导药(Solution)的展示方式
- 从"购买顾虑"推导信(Trust)需要解决的信任问题
- 从"情绪触发词"提炼口播文案的关键词
- 从"使用场景"推导视频的画面和情境设计

===== 任务要求 =====

请为产品【{product}】生成 **6个不同角度** 的短视频广告脚本。

每个角度的脚本格式：

### 角度#N: [角度名称] — [使用的结构公式名称]

**目标受众：** 谁会被这条视频吸引
**视频时长：** 建议时长
**适合平台：** TikTok / Instagram Reels / Facebook Reels

**停(Hook) [前3秒]：**
- 画面描述: 具体到镜头运动和产品动作
- 口播/字幕: "..." (英文)
- 钩子类型: 痛点冲击/效果冲击/好奇悬念/对比冲击
- 来源: [素材库] 引用ID / [AI原创] 基于调研

**病(Pain) [3-8秒]：**
- 画面描述: 具体到场景和演员动作
- 口播/字幕: "..." (英文)
- 放大的痛点: 用一句话概括
- 来源: [素材库] / [AI原创]

**药(Solution) [8-20秒]：**
- 画面描述: 具体到产品的哪个卖点、怎么拍、什么角度
- 口播/字幕: "..." (英文)

**信(Trust) [20-25秒]：**
- 画面描述: 用什么社会证明（好评/认证/销量/达人）
- 口播/字幕: "..." (英文)

**买(CTA) [最后5秒]：**
- 画面描述: 具体到促销信息的展示方式
- 口播/字幕: "..." (英文)
- 促单权益: 折扣/包邮/限时等

**漏斗覆盖：** TOFU/MOFU/BOFU（标注每层如何覆盖）

---

要求：
1. 6个角度必须切入不同的痛点/情绪/场景，不要重复
2. 至少使用3种不同的结构公式（从12种中选择最合适的）
3. 口播文案全部用英文（美区市场），画面描述用中文
4. 参考竞品的有效策略，但要做出差异化
5. 每条脚本必须覆盖至少2层漏斗，其中至少2条要覆盖全部3层（TOFU+MOFU+BOFU）
6. 结合素材库已有的痛点和评论，但也要提出新角度
'''

    try:
        ai_text = deepseek(system_prompt, user_prompt)
        log(f'  AI: {len(ai_text)} chars')
    except Exception as e:
        log(f'  DeepSeek failed: {e}')
        return

    # ── Step 4: Build .docx ───────────────────────────────────
    log('Step 4: Building .docx...')
    report_title = f'({date_str})VSL脚本-{product}'
    docx_path = build_vsl_docx(report_title, product, features, ai_text, now)
    log(f'  DOCX: {os.path.getsize(docx_path)} bytes')

    # ── Step 5: Import to wiki ────────────────────────────────
    log('Step 5: Importing to Feishu...')
    wiki_url, doc_token = import_docx_to_wiki(docx_path, report_title)
    try:
        os.remove(docx_path)
    except OSError:
        pass
    if not wiki_url:
        log('  Failed!')
        return
    log(f'  Report: {wiki_url}')

    # ── Step 6: Notify ────────────────────────────────────────
    log('Step 6: Sending notification...')
    card = {
        'config': {'wide_screen_mode': True},
        'header': {
            'title': {'tag': 'plain_text', 'content': f'VSL脚本生成 — {product}'},
            'template': 'green',
        },
        'elements': [
            {'tag': 'div', 'text': {'tag': 'lark_md',
                'content': f'**产品:** {product}\n**卖点:** {features}\n**角度:** 6个VSL脚本已生成'}},
            {'tag': 'hr'},
            {'tag': 'action', 'actions': [{
                'tag': 'button',
                'text': {'tag': 'plain_text', 'content': '查看脚本'},
                'url': wiki_url, 'type': 'primary',
            }]},
        ],
    }
    msg_body = json.dumps(card)
    feishu('POST', '/im/v1/messages?receive_id_type=open_id',
           {'receive_id': BOSS_OPEN_ID, 'msg_type': 'interactive', 'content': msg_body}, app='im')
    log('  Notification sent')

    # ── Step 7: Parse scripts into JSON and split into material libraries ──
    log('Step 7: Parsing scripts into structured data...')
    total_records = split_scripts_to_bitable(ai_text, product, now)
    log(f'  Total records created: {total_records}')

    log(f'\nS4 Complete! {wiki_url}')

# ============================================================
# Split scripts into material libraries
# ============================================================
FRAMEWORK_MAP = {
    '经典痛点': 'recvdOMR9Zpm7g',
    '效果前置': 'recvdOMWKAWDVh',
    '对比碾压': 'recvdOMXyQtizB',
    '好奇悬念': 'recvdOMZI8brzb',
    '社交证明': 'recvdON1Evlm6V',
    '真实体验': 'recvdON3ZrLG5O',
    '科普权威': 'recvdON5QdXzJ5',
    '多场景轰炸': 'recvdON7EBmesm',
    '开箱种草': 'recvdON9J95f4Z',
    '剧情反转': 'recvdONbTT1QHu',
    '纯产品展示': 'recveBGraDfJsO',
    '极短展示': 'recveWNjzBq8uG',
}

def match_framework(title):
    """Match angle title to framework record_id."""
    for key, rid in FRAMEWORK_MAP.items():
        if key in title:
            return rid
    return None

def create_record(table_id, fields):
    """Create a bitable record, return record_id."""
    r = feishu('POST', f'/bitable/v1/apps/{MATERIAL_APP}/tables/{table_id}/records', {'fields': fields})
    if r.get('code') == 0:
        return r['data']['record']['record_id']
    else:
        log(f'    Write failed: {str(r)[:150]}')
        return None

def split_scripts_to_bitable(ai_text, product, now):
    """Parse AI-generated scripts and write to material libraries + task table."""
    # Use DeepSeek to parse scripts into structured JSON
    parse_prompt = '''请将以下VSL脚本解析为JSON数组。每个角度提取以下字段：

```json
[
  {
    "angle_name": "角度名称",
    "framework": "使用的结构公式名称（只写关键词如：经典痛点/效果前置/对比碾压/好奇悬念/社交证明/真实体验/科普权威/多场景轰炸/开箱种草/剧情反转/纯产品展示/极短展示）",
    "audience": "目标受众",
    "duration": "视频时长",
    "hook_text_en": "停的英文口播/字幕",
    "hook_type": "钩子类型",
    "hook_scene": "停的画面描述",
    "pain_text_en": "病的英文口播/字幕",
    "pain_point": "放大的痛点（一句话）",
    "pain_scene": "病的画面描述",
    "pain_emotions": ["情绪1", "情绪2"],
    "solution_text_en": "药的英文口播/字幕",
    "solution_scene": "药的画面描述",
    "trust_items": [
      {"type": "权威认证", "content": "具体内容", "text_en": "英文话术"},
      {"type": "用户好评", "content": "具体内容", "text_en": "英文话术"}
    ],
    "cta_text_en": "买的英文口播/字幕",
    "cta_benefit": "促单权益（如15%折扣）",
    "cta_scene": "买的画面描述",
    "funnel": "漏斗覆盖"
  }
]
```

**重要规则：**
1. pain_emotions 必须从以下选项中选2-3个：不满/焦虑/遗憾/尴尬/渴望/失望/无奈/惊喜/好奇/恐惧/愧疚。每个角度的情绪组合应不同。
2. trust_items 必须拆分为独立条目，每种类型单独一条。如果脚本中同时提到了"官方认证"和"用户好评"，就输出2个对象。type只能是：权威认证/用户好评/销量数据/网红背书。
3. 每个角度至少输出2条trust_items。

只输出JSON，不要其他文字。
脚本内容：
''' + ai_text

    try:
        json_text = deepseek('你是JSON解析专家，只输出合法JSON。', parse_prompt)
        # Extract JSON from response
        json_match = re.search(r'\[[\s\S]*\]', json_text)
        if json_match:
            angles = json.loads(json_match.group())
        else:
            log('  Failed to parse JSON')
            return 0
        log(f'  Parsed {len(angles)} angles into JSON')
    except Exception as e:
        log(f'  Parse failed: {e}')
        return 0

    # Calculate this week's Monday
    today = now.date()
    import datetime as _dt
    this_monday = today - _dt.timedelta(days=today.weekday())
    monday_ts = int(_dt.datetime.combine(this_monday, _dt.time()).timestamp() * 1000)

    total = 0
    for i, angle in enumerate(angles):
        log(f'  Angle #{i+1}: {angle.get("angle_name","")}')
        linked_ids = {'cta': [], 'pain': [], 'solution': [], 'trust': [], 'benefit': []}

        # 1. CTA库 — 开头钩子
        hook_text = angle.get('hook_text_en', '')
        if hook_text:
            rid = create_record('tblbtMGxaYTAUqLk', {
                'CTA话术（英文）': hook_text,
                '中文翻译': angle.get('hook_scene', '')[:200],
                '视频阶段': '开头（前3秒钩子）',
                '话术逻辑': f'钩子类型: {angle.get("hook_type","")}',
                '配合画面': angle.get('hook_scene', '')[:200],
                '适用品类': ['Switch配件'],
            })
            if rid:
                linked_ids['cta'].append(rid)
                total += 1

        # 2. 痛点需求场景库
        pain_point = angle.get('pain_point', '')
        if pain_point:
            # Use emotions from DeepSeek JSON (not code inference)
            emotion_words = angle.get('pain_emotions', ['不满'])
            if not emotion_words:
                emotion_words = ['不满']

            rid = create_record('tblxT2UNUaI89Pjj', {
                '场景名称': angle.get('angle_name', '')[:50],
                '场景分类': '✨ 产品体验',
                '用户痛点': pain_point,
                '情绪关键词': emotion_words,  # multi-select, individual items
                '产品切入点': product,
                '内容角度建议': angle.get('pain_scene', '')[:200],
                '话术示例（英文）': angle.get('pain_text_en', ''),
                '来源': 'AI生成',
                '适用品类': ['Switch配件'],
            })
            if rid:
                linked_ids['pain'].append(rid)
                total += 1

        # 3. 卖点画面库
        solution_scene = angle.get('solution_scene', '')
        if solution_scene:
            # Detect scene type for 画面类型
            scene_type = '产品展示'
            scene_lower = solution_scene.lower()
            if any(kw in scene_lower for kw in ['对比', '分屏', 'vs']):
                scene_type = 'Before/After对比'
            elif any(kw in scene_lower for kw in ['开箱', '拆', '包装']):
                scene_type = '开箱展示'
            elif any(kw in scene_lower for kw in ['使用', '操作', '实际']):
                scene_type = '使用演示'

            # If ASMR mentioned, note in 作用 field
            action_text = angle.get('solution_text_en', '')[:200]
            if 'asmr' in scene_lower or 'ASMR' in solution_scene:
                action_text += '（配合ASMR声效增强沉浸感）'

            rid = create_record('tblIkl7J7R3T9QgX', {
                '拍摄说明': solution_scene[:500],
                '画面类型': scene_type,
                '作用': action_text,
                '配合视频阶段': '中间（卖点引导）',
                '适用品类': ['Switch配件'],
            })
            if rid:
                linked_ids['solution'].append(rid)
                total += 1

        # 4. 社会证明库 — 逐条拆分
        for trust_item in angle.get('trust_items', []):
            t_type = trust_item.get('type', '用户好评')
            # Map to valid options
            valid_types = {'权威认证': '权威认证', '用户好评': '用户好评',
                          '销量数据': '销量数据', '网红背书': '网红背书',
                          '官方认证': '权威认证'}
            mapped_type = valid_types.get(t_type, '用户好评')

            rid = create_record('tbl3lI2MTHkbnWHR', {
                '素材名称': trust_item.get('content', '')[:100],
                '证明类型': mapped_type,
                '素材来源': 'AI生成',
                '使用场景说明': trust_item.get('text_en', '')[:200],
                '信任强度': '强',
            })
            if rid:
                linked_ids['trust'].append(rid)
                total += 1

        # 5. 权益库
        cta_benefit = angle.get('cta_benefit', '')
        if cta_benefit:
            rid = create_record('tblI1RssOVWfVIgw', {
                '权益名称': cta_benefit[:50],
                '权益类型': '限时折扣' if '折扣' in cta_benefit or 'OFF' in cta_benefit else '其他',
                '权益描述': angle.get('cta_scene', '')[:200],
                '话术示例': angle.get('cta_text_en', '')[:200],
                '适用品类': ['Switch配件'],
            })
            if rid:
                linked_ids['benefit'].append(rid)
                total += 1

        # 6. CTA库 — 结尾CTA
        cta_text = angle.get('cta_text_en', '')
        if cta_text:
            rid = create_record('tblbtMGxaYTAUqLk', {
                'CTA话术（英文）': cta_text,
                '中文翻译': angle.get('cta_scene', '')[:200],
                '视频阶段': '结尾（促单转化）',
                '话术逻辑': f'促单权益: {cta_benefit}',
                '配合权益': cta_benefit,
                '适用品类': ['Switch配件'],
            })
            if rid:
                linked_ids['cta'].append(rid)
                total += 1

        # 7. 每周视频任务需求
        framework_rid = match_framework(angle.get('framework', ''))
        task_fields = {
            '时间': monday_ts,
            '产品': product,
            '发布平台': 'FB广告',
            '语种': '英语',
            '任务状态': '未完成',
        }
        if framework_rid:
            task_fields['结构公式'] = [framework_rid]
        if linked_ids['cta']:
            task_fields['停·素材来源'] = linked_ids['cta']
        if linked_ids['pain']:
            task_fields['病·素材来源'] = linked_ids['pain']
        if linked_ids['solution']:
            task_fields['药·素材来源'] = linked_ids['solution']
        if linked_ids['trust']:
            task_fields['信·素材来源'] = linked_ids['trust']
        if linked_ids['benefit'] or linked_ids['cta']:
            task_fields['买·素材来源'] = linked_ids['benefit'] + linked_ids['cta'][-1:]

        rid = create_record('tbl3tATWMmLm60SG', task_fields)
        if rid:
            total += 1
            log(f'    Task created: {rid}')

    return total

if __name__ == '__main__':
    main()
