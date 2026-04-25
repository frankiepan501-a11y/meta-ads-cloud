"""S1 Meta Ads Weekly Report - Production Script.
Pulls Facebook Ads data → DeepSeek analysis → python-docx → import to Feishu wiki.
Triggered via FastAPI POST /run/s1 (cron orchestrated by n8n).
"""
import json, urllib.request, sys, socket, ssl, re, datetime, time, uuid, os, io
try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

from cloud_config import (
    PROXY, IM_APP_ID, IM_APP_SECRET, DEEPSEEK_KEY,
    PK_TOKEN, FL_TOKEN, PK_ACCOUNT, FL_ACCOUNT,
    WIKI_SPACE_ID, BOSS_OPEN_ID, SKIP_ADLIB_IMAGES,
)
WIKI_PARENT_NODE = os.environ.get('WIKI_PARENT_NODE_S1', 'DLlkwhBymiHJfjk5mC6cuqNnnDy')

LOG_FILE = os.environ.get('S1_LOG_FILE', '/tmp/s1_log.txt')

def log(msg):
    ts = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    line = f'[{ts}] {msg}'
    print(line)
    with open(LOG_FILE, 'a', encoding='utf-8') as f:
        f.write(line + '\n')

# ============================================================
# API helpers
# ============================================================
_token_cache = {}

def get_token(app_id, secret):
    key = app_id
    if key in _token_cache and time.time() - _token_cache[key][1] < 600:
        return _token_cache[key][0]
    body = json.dumps({'app_id': app_id, 'app_secret': secret}).encode()
    req = urllib.request.Request('https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal',
        data=body, headers={'Content-Type': 'application/json'})
    token = json.loads(urllib.request.build_opener(PROXY).open(req, timeout=15).read())['tenant_access_token']
    _token_cache[key] = (token, time.time())
    return token

def feishu(method, path, body=None):
    url = f'https://open.feishu.cn/open-apis{path}'
    data = json.dumps(body).encode('utf-8') if body else None
    last_err = None
    for attempt in range(3):
        token = get_token(IM_APP_ID, IM_APP_SECRET)
        req = urllib.request.Request(url, data=data,
            headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {token}'}, method=method)
        try:
            return json.loads(urllib.request.build_opener(PROXY).open(req, timeout=30).read())
        except urllib.error.HTTPError as e:
            err = e.read().decode('utf-8', 'replace')[:500]
            log(f'  API Error {e.code}: {err[:200]}')
            return {'code': e.code, 'body': err}
        except (urllib.error.URLError, OSError) as e:
            last_err = e
            log(f'  Network error (attempt {attempt+1}/3): {e}')
            time.sleep(2 * (attempt + 1))
    return {'code': -1, 'body': f'Network error after 3 retries: {last_err}'}

def get_users_by_job_title(title):
    r = feishu('GET', '/contact/v3/departments?parent_department_id=0&page_size=50'
               '&department_id_type=open_department_id&user_id_type=open_id&fetch_child=true')
    dept_ids = ['0'] + [d['open_department_id'] for d in r.get('data', {}).get('items', [])]
    seen, matched = set(), []
    for did in dept_ids:
        page = ''
        while True:
            q = (f'/contact/v3/users/find_by_department?department_id={did}'
                 f'&page_size=50&user_id_type=open_id&department_id_type=open_department_id')
            if page:
                q += f'&page_token={page}'
            r = feishu('GET', q)
            for u in r.get('data', {}).get('items', []):
                oid = u.get('open_id')
                if not oid or oid in seen:
                    continue
                seen.add(oid)
                if u.get('job_title') == title:
                    matched.append((u.get('name'), oid))
            page = r.get('data', {}).get('page_token', '')
            if not page:
                break
    return matched

def fb_api(path, token):
    """Call Facebook Graph API via urllib (PROXY honors HTTP_PROXY env)."""
    full_path = f'{path}&access_token={token}' if '?' in path else f'{path}?access_token={token}'
    url = 'https://graph.facebook.com' + full_path
    req = urllib.request.Request(url, headers={'User-Agent': 'meta-ads-cloud/1.0'})
    last_err = None
    for attempt in range(3):
        try:
            return json.loads(urllib.request.build_opener(PROXY).open(req, timeout=45).read())
        except urllib.error.HTTPError as e:
            err = e.read().decode('utf-8', 'replace')[:500]
            log(f'  FB API Error {e.code}: {err[:200]}')
            return {'error': {'code': e.code, 'message': err}}
        except (urllib.error.URLError, OSError) as e:
            last_err = e
            log(f'  FB network error (attempt {attempt+1}/3): {e}')
            time.sleep(2 * (attempt + 1))
    return {'error': {'message': f'network error after 3 retries: {last_err}'}}

def deepseek(system_prompt, user_prompt):
    body = json.dumps({
        'model': 'deepseek-chat', 'max_tokens': 6000,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt},
        ]
    }).encode('utf-8')
    req = urllib.request.Request('https://api.deepseek.com/chat/completions',
        data=body, headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {DEEPSEEK_KEY}'})
    resp = json.loads(urllib.request.build_opener(PROXY).open(req, timeout=120).read())
    return resp['choices'][0]['message']['content']

def download_image(url):
    """Download image from URL, return bytes or None."""
    try:
        req = urllib.request.Request(url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'})
        data = urllib.request.build_opener(PROXY).open(req, timeout=15).read()
        return data if len(data) > 500 else None
    except Exception:
        return None

# ============================================================
# Docx import to Feishu wiki
# ============================================================
def import_docx_to_wiki(docx_path, title):
    """Upload .docx, import as Feishu doc, move to wiki. Returns (wiki_url, doc_token)."""
    token = get_token(IM_APP_ID, IM_APP_SECRET)
    with open(docx_path, 'rb') as f:
        docx_bytes = f.read()

    boundary = uuid.uuid4().hex
    safe_name = f'{title}.docx'
    extra_json = json.dumps({'obj_type': 'docx', 'file_extension': 'docx'})
    parts = []
    for name, value in [('file_name', safe_name), ('parent_type', 'ccm_import_open'),
                        ('size', str(len(docx_bytes))), ('extra', extra_json)]:
        parts.append(f'--{boundary}\r\nContent-Disposition: form-data; name="{name}"\r\n\r\n{value}\r\n'.encode())
    parts.append(f'--{boundary}\r\nContent-Disposition: form-data; name="file"; filename="{safe_name}"\r\n'
                 f'Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n'.encode())
    parts.append(docx_bytes)
    parts.append(f'\r\n--{boundary}--\r\n'.encode())

    req = urllib.request.Request('https://open.feishu.cn/open-apis/drive/v1/medias/upload_all',
        data=b''.join(parts),
        headers={'Content-Type': f'multipart/form-data; boundary={boundary}', 'Authorization': f'Bearer {token}'})
    try:
        resp = json.loads(urllib.request.build_opener(PROXY).open(req, timeout=60).read())
        file_token = resp['data']['file_token']
    except Exception as e:
        log(f'  DOCX upload failed: {e}')
        return None, None

    r = feishu('POST', '/drive/v1/import_tasks', {
        'file_extension': 'docx', 'file_token': file_token,
        'type': 'docx', 'point': {'mount_type': 1, 'mount_key': ''},
    })
    if r.get('code') != 0:
        log(f'  Import task failed: {r}')
        return None, None
    ticket = r['data']['ticket']

    doc_token = None
    for _ in range(60):
        time.sleep(2)
        r2 = feishu('GET', f'/drive/v1/import_tasks/{ticket}')
        status = r2.get('data', {}).get('result', {}).get('job_status', -1)
        if status == 0:
            doc_token = r2['data']['result'].get('token', '')
            break
        elif status in (1, 2):
            log(f'  Import failed: {r2["data"]["result"].get("job_error_msg","")}')
            return None, None
    if not doc_token:
        log('  Import timeout')
        return None, None

    feishu('POST', f'/drive/v1/permissions/{doc_token}/members?type=docx&need_notification=false',
           {'member_type': 'openid', 'member_id': BOSS_OPEN_ID, 'perm': 'full_access'})

    r3 = feishu('POST', f'/wiki/v2/spaces/{WIKI_SPACE_ID}/nodes/move_docs_to_wiki', {
        'parent_wiki_token': WIKI_PARENT_NODE,
        'obj_type': 'docx', 'obj_token': doc_token,
    })
    if r3.get('code') != 0:
        log(f'  Move to wiki failed: {r3}')
        return None, None

    node_token = r3.get('data', {}).get('wiki_token')
    if not node_token:
        for _ in range(10):
            time.sleep(2)
            rn = feishu('GET', f'/wiki/v2/spaces/get_node?obj_type=docx&token={doc_token}')
            node_token = rn.get('data', {}).get('node', {}).get('node_token')
            if node_token:
                break
        if not node_token:
            log(f'  Wiki node_token not resolved: {r3}')
            return None, None

    wiki_url = f'https://u1wpma3xuhr.feishu.cn/wiki/{node_token}'
    return wiki_url, doc_token

# ============================================================
# Facebook data helpers
# ============================================================
def get_insights(account_id, token, time_range, level='campaign'):
    import urllib.parse
    tr = urllib.parse.quote(json.dumps(time_range))
    fields = 'campaign_id,campaign_name,spend,impressions,clicks,ctr,cpc,cpm,actions,purchase_roas'
    if level == 'ad':
        fields = 'ad_id,ad_name,adset_name,campaign_name,spend,impressions,clicks,ctr,cpc,actions,purchase_roas'
    path = f'/v21.0/{account_id}/insights?fields={fields}&time_range={tr}&level={level}&limit=50'
    if level == 'ad':
        path += '&sort=spend_descending'
    return fb_api(path, token)

def get_ad_creatives(account_id, token):
    """Get active ads with creative thumbnail. Returns {ad_id: creative_data}."""
    import urllib.parse
    es = urllib.parse.quote('["ACTIVE"]')
    path = (f'/v21.0/{account_id}/ads?fields=id,name,creative'
            f'{{id,title,body,thumbnail_url,effective_object_story_id}}'
            f'&effective_status={es}&limit=30')
    r = fb_api(path, token)
    result = {}
    for ad in r.get('data', []):
        c = ad.get('creative', {})
        result[ad['id']] = {
            'body': c.get('body', ''),
            'thumbnail_url': c.get('thumbnail_url', ''),
            'post_id': c.get('effective_object_story_id', ''),
        }
    return result

def scrape_own_ads(brand_name, page_names):
    """Scrape Ad Library for own brand's HD creative images.
    Returns list of {pageName, body, images, hasVideo, mediaType, id}.
    v1: skipped on cloud — Ad Library API approval pending. Returns [] when SKIP_ADLIB_IMAGES.
    """
    if SKIP_ADLIB_IMAGES:
        log(f'  Ad Library scrape skipped for {brand_name} (SKIP_ADLIB_IMAGES=1)')
        return []
    from fb_ad_library_scraper import scrape_ad_library
    try:
        all_ads = scrape_ad_library(brand_name, 'ALL', 20)
        pn_lower = [p.lower() for p in page_names]
        own = [a for a in all_ads if a.get('pageName', '').lower() in pn_lower]
        return own
    except Exception as e:
        log(f'  Ad Library scrape failed for {brand_name}: {e}')
        return []

def parse_action(actions, atype):
    if not actions: return 0
    for a in actions:
        if a.get('action_type') == atype: return float(a.get('value', 0))
    return 0

def parse_roas(roas_arr):
    if not roas_arr: return 0
    for r in roas_arr:
        if r.get('action_type') in ('omni_purchase', 'purchase'):
            return float(r.get('value', 0))
    return float(roas_arr[0].get('value', 0)) if roas_arr else 0

def sum_account(data):
    total = {'spend': 0, 'impressions': 0, 'clicks': 0, 'purchases': 0, 'revenue': 0}
    campaigns = []
    for d in data.get('data', []):
        spend = float(d.get('spend', 0))
        roas = parse_roas(d.get('purchase_roas'))
        purchases = parse_action(d.get('actions'), 'purchase') or parse_action(d.get('actions'), 'omni_purchase')
        total['spend'] += spend
        total['impressions'] += int(d.get('impressions', 0))
        total['clicks'] += int(d.get('clicks', 0))
        total['purchases'] += purchases
        total['revenue'] += spend * roas
        campaigns.append({
            'name': d.get('campaign_name', ''),
            'spend': round(spend, 2),
            'roas': round(roas, 2),
            'purchases': int(purchases),
            'cpc': round(float(d.get('cpc', 0)), 2),
        })
    campaigns.sort(key=lambda x: x['roas'], reverse=True)
    total['roas'] = round(total['revenue'] / max(total['spend'], 0.01), 2)
    total['cpa'] = round(total['spend'] / max(total['purchases'], 1), 2)
    return total, campaigns

def parse_top_ads(insights_data, creatives_map, account):
    """Merge ad insights with creative data.
    creatives_map: {ad_id: {body, thumbnail_url, post_id}}
    """
    ads = []
    for d in insights_data.get('data', []):
        spend = float(d.get('spend', 0))
        roas = parse_roas(d.get('purchase_roas'))
        purchases = parse_action(d.get('actions'), 'purchase') or parse_action(d.get('actions'), 'omni_purchase')
        ad_id = d.get('ad_id', '')
        ad_name = d.get('ad_name', '')

        creative = creatives_map.get(ad_id, {})
        post_id = creative.get('post_id', '')
        creative_type = 'video' if any(kw in ad_name.lower() for kw in ['视频', 'video', 'ugc']) else 'image'

        ads.append({
            'account': account,
            'ad_id': ad_id,
            'ad_name': ad_name,
            'campaign': d.get('campaign_name', ''),
            'adset': d.get('adset_name', ''),
            'spend': round(spend, 2),
            'roas': round(roas, 2),
            'purchases': int(purchases),
            'cpc': round(float(d.get('cpc', 0)), 2),
            'ctr': round(float(d.get('ctr', 0)), 2),
            'creative_type': creative_type,
            'best_image': creative.get('thumbnail_url', ''),
            'body': creative.get('body', ''),
            'fb_url': f'https://www.facebook.com/{post_id}' if post_id else '',
        })
    ads.sort(key=lambda x: x['roas'], reverse=True)
    return ads

def wow(curr, prev):
    c, p = float(curr), float(prev)
    if p == 0: return '+100%' if c > 0 else '0%'
    pct = round((c - p) / p * 100, 1)
    return f'+{pct}%' if pct > 0 else f'{pct}%'

# ============================================================
# Build .docx report
# ============================================================
def build_report_docx(title, pk_total, pk_last_total, fl_total, fl_last_total,
                      pk_campaigns, fl_campaigns, all_top_ads, sections,
                      ad_library_ads, now):
    """Build .docx with embedded ad images from Ad Library."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    import tempfile

    doc = Document()

    # Title
    doc.add_heading(title, level=1)
    doc.add_paragraph(f'生成时间: {now.strftime("%Y-%m-%d %H:%M")}')

    # ── Overview ──
    doc.add_heading('整体表现', level=2)

    wow_pk_spend = wow(pk_total['spend'], pk_last_total['spend'])
    wow_pk_roas = wow(pk_total['roas'], pk_last_total['roas'])
    wow_fl_spend = wow(fl_total['spend'], fl_last_total['spend'])
    wow_fl_roas = wow(fl_total['roas'], fl_last_total['roas'])

    p = doc.add_paragraph()
    p.add_run('Powkong | ').bold = True
    p.add_run(f'Spend: ${pk_total["spend"]:.2f} ({wow_pk_spend}) | '
              f'ROAS: {pk_total["roas"]} ({wow_pk_roas}) | '
              f'Purchases: {pk_total["purchases"]}')

    p = doc.add_paragraph()
    p.add_run('Funlab   | ').bold = True
    p.add_run(f'Spend: ${fl_total["spend"]:.2f} ({wow_fl_spend}) | '
              f'ROAS: {fl_total["roas"]} ({wow_fl_roas}) | '
              f'Purchases: {fl_total["purchases"]}')

    # ── AI sections (text-based) ──
    for key, title_text in [('整体表现总结', 'AI 分析总结'),
                            ('Campaign排行', 'Campaign 排行'),
                            ('问题诊断', '问题诊断')]:
        content = sections.get(key, '').strip()
        if not content:
            continue
        doc.add_heading(title_text, level=2)
        for line in content.split('\n')[:25]:
            t = line.strip()
            if not t: continue
            if t.startswith(('- ', '* ', '• ')):
                doc.add_paragraph(re.sub(r'^[-*•]\s*', '', t), style='List Bullet')
            elif re.match(r'^\d+[\.\)、]', t):
                doc.add_paragraph(re.sub(r'^\d+[\.\)、]\s*', '', t), style='List Number')
            else:
                doc.add_paragraph(t)

    # ── Winning Ads: 图片 + 数据 + AI拆解 合并展示 ──
    doc.add_heading('Winning Ads 素材展示', level=2)

    # Build Ad Library image lookup
    lib_ads_with_img = [a for a in ad_library_ads if a.get('images')]
    lib_used_idx = set()

    def find_lib_image(ad_body, ad_name):
        """Find best matching Ad Library HD image."""
        if ad_body:
            prefix = ad_body[:40].strip()
            for idx, la in enumerate(lib_ads_with_img):
                if idx not in lib_used_idx and prefix and prefix in la.get('body', ''):
                    lib_used_idx.add(idx)
                    return la['images'][0]
        keywords = [kw for kw in ad_name.split() if len(kw) > 1]
        for idx, la in enumerate(lib_ads_with_img):
            if idx not in lib_used_idx:
                if any(kw.lower() in la.get('body', '').lower() for kw in keywords if len(kw) > 2):
                    lib_used_idx.add(idx)
                    return la['images'][0]
        for idx, la in enumerate(lib_ads_with_img):
            if idx not in lib_used_idx:
                lib_used_idx.add(idx)
                return la['images'][0]
        return ''

    # Parse AI per-ad analysis into chunks
    winning_text = sections.get('Winning Ads素材分析', '').strip()
    ad_analyses = re.split(r'\n(?=\d+[\.\)、])', winning_text) if winning_text else []

    img_embedded = 0
    for i, ad in enumerate(all_top_ads[:8]):
        winning = ' ✅' if ad['roas'] >= 2.0 else ''
        type_label = '🎬 视频' if ad['creative_type'] == 'video' else '🖼️ 图片'

        doc.add_heading(f'#{i+1} {ad["ad_name"]}{winning}', level=3)

        # HD image from Ad Library
        img_url = find_lib_image(ad.get('body', ''), ad.get('ad_name', ''))
        if img_url:
            img_data = download_image(img_url)
            if img_data:
                try:
                    doc.add_picture(io.BytesIO(img_data), width=Inches(4.5))
                    img_embedded += 1
                except Exception:
                    pass

        # Performance metrics
        p = doc.add_paragraph()
        p.add_run('ROAS: ').bold = True
        p.add_run(f'{ad["roas"]} | ')
        p.add_run('Spend: ').bold = True
        p.add_run(f'${ad["spend"]} | ')
        p.add_run('Purchases: ').bold = True
        p.add_run(f'{ad["purchases"]} | ')
        p.add_run('CPC: ').bold = True
        p.add_run(f'${ad["cpc"]} | ')
        p.add_run('CTR: ').bold = True
        p.add_run(f'{ad["ctr"]}%')

        p = doc.add_paragraph()
        p.add_run('素材类型: ').bold = True
        p.add_run(f'{type_label} | ')
        p.add_run('Account: ').bold = True
        p.add_run(f'{ad["account"]} | ')
        p.add_run('Campaign: ').bold = True
        p.add_run(ad['campaign'])

        # AI analysis for this ad (if available)
        if i < len(ad_analyses):
            for line in ad_analyses[i].split('\n'):
                t = line.strip()
                if not t:
                    continue
                if t.startswith(('- ', '* ', '• ')):
                    doc.add_paragraph(re.sub(r'^[-*•]\s*', '', t), style='List Bullet')
                elif not re.match(r'^\d+[\.\)、]', t):
                    doc.add_paragraph(t, style='List Bullet')

        if ad['fb_url']:
            doc.add_paragraph(f'🔗 查看原始广告帖: {ad["fb_url"]}')

    log(f'  Ad Library images embedded: {img_embedded}')

    # ── More AI sections ──
    for key, title_text in [('本周假设验证', '本周假设验证'),
                            ('下周内容制作指引', '下周内容制作指引')]:
        content = sections.get(key, '').strip()
        if not content: continue
        doc.add_heading(title_text, level=2)
        for line in content.split('\n')[:20]:
            t = line.strip()
            if not t: continue
            if t.startswith(('- ', '* ', '• ')):
                doc.add_paragraph(re.sub(r'^[-*•]\s*', '', t), style='List Bullet')
            elif re.match(r'^\d+[\.\)、]', t):
                doc.add_paragraph(re.sub(r'^\d+[\.\)、]\s*', '', t), style='List Number')
            else:
                doc.add_paragraph(t)

    # ── Action items ──
    actions = sections.get('下周行动项', '').strip()
    if actions:
        doc.add_heading('下周行动项', level=2)
        for line in actions.split('\n')[:10]:
            t = line.strip()
            if not t: continue
            clean = re.sub(r'^[\d\.\)、\-*•]+\s*', '', t)
            clean = re.sub(r'^\[?\s*\]?\s*', '', clean)
            if clean:
                doc.add_paragraph(f'☐ {clean}', style='List Bullet')

    path = os.path.join(os.environ.get('TEMP', '/tmp'), 's1_report.docx')
    doc.save(path)
    return path

# ============================================================
# Main
# ============================================================
def main():
    log('=' * 60)
    log('S1 Meta Ads Weekly Report')

    now = datetime.datetime.now()
    fake_today = os.environ.get('META_FAKE_TODAY')
    if fake_today:
        now = datetime.datetime.strptime(fake_today, '%Y-%m-%d')
        log(f'  [OVERRIDE] today = {fake_today}')
    date_str = now.strftime('%Y-%m-%d')

    # Auto-calculate this week (Mon-Sun ending yesterday) and last week
    today = now.date()
    # This week = last Monday to last Sunday
    days_since_monday = today.weekday()  # 0=Mon
    this_sun = today - datetime.timedelta(days=max(days_since_monday, 1))
    this_mon = this_sun - datetime.timedelta(days=6)
    last_sun = this_mon - datetime.timedelta(days=1)
    last_mon = last_sun - datetime.timedelta(days=6)

    this_week = {'since': this_mon.isoformat(), 'until': this_sun.isoformat()}
    last_week = {'since': last_mon.isoformat(), 'until': last_sun.isoformat()}
    week_num = this_mon.isocalendar()[1]
    log(f'  This week: {this_week["since"]} ~ {this_week["until"]} (W{week_num})')
    log(f'  Last week: {last_week["since"]} ~ {last_week["until"]}')

    # ── Step 1: Pull Facebook data ────────────────────────────
    log('Step 1: Pulling Facebook data...')
    pk_this = get_insights(PK_ACCOUNT, PK_TOKEN, this_week)
    pk_last = get_insights(PK_ACCOUNT, PK_TOKEN, last_week)
    pk_ads = get_insights(PK_ACCOUNT, PK_TOKEN, this_week, 'ad')
    pk_creatives = get_ad_creatives(PK_ACCOUNT, PK_TOKEN)

    fl_this = get_insights(FL_ACCOUNT, FL_TOKEN, this_week)
    fl_last = get_insights(FL_ACCOUNT, FL_TOKEN, last_week)
    fl_ads = get_insights(FL_ACCOUNT, FL_TOKEN, this_week, 'ad')
    fl_creatives = get_ad_creatives(FL_ACCOUNT, FL_TOKEN)

    log(f'  PK: {len(pk_this.get("data",[]))} campaigns, {len(pk_ads.get("data",[]))} ads, {len(pk_creatives)} creatives')
    log(f'  FL: {len(fl_this.get("data",[]))} campaigns, {len(fl_ads.get("data",[]))} ads, {len(fl_creatives)} creatives')

    if pk_this.get('error') or fl_this.get('error'):
        log(f'  Facebook API error: PK={pk_this.get("error")} FL={fl_this.get("error")}')
        log('  Aborting.')
        return

    # ── Step 2: Parse data ────────────────────────────────────
    log('Step 2: Parsing data...')
    pk_total, pk_campaigns = sum_account(pk_this)
    pk_last_total, _ = sum_account(pk_last)
    fl_total, fl_campaigns = sum_account(fl_this)
    fl_last_total, _ = sum_account(fl_last)

    pk_top_ads = parse_top_ads(pk_ads, pk_creatives, 'Powkong')
    fl_top_ads = parse_top_ads(fl_ads, fl_creatives, 'Funlab')
    all_top_ads = sorted(pk_top_ads + fl_top_ads, key=lambda x: x['roas'], reverse=True)

    log(f'  Top ads: {len(all_top_ads)} (with images: {sum(1 for a in all_top_ads if a.get("best_image"))})')

    # ── Step 3: Scrape Ad Library for HD creative images ─────
    log('Step 3: Scraping Ad Library for HD images...')
    pk_lib = scrape_own_ads('Powkong', ['Powkong'])
    fl_lib = scrape_own_ads('FUNLABOFFICIAL', ['FUNLABOFFICIAL'])
    ad_library_ads = pk_lib + fl_lib
    log(f'  Ad Library: Powkong={len(pk_lib)} Funlab={len(fl_lib)} '
        f'(with images: {sum(1 for a in ad_library_ads if a.get("images"))})')

    # ── Step 4: DeepSeek analysis ─────────────────────────────
    log('Step 4: Calling DeepSeek...')

    kb_path = os.environ.get('META_ADS_KB_PATH', os.path.join(os.path.dirname(os.path.abspath(__file__)), 'meta_ads_knowledge.md'))
    with open(kb_path, 'r', encoding='utf-8') as f:
        system_prompt = f.read()
    system_prompt += '\n\n## 品牌信息\n- Powkong（宝空）：游戏配件品牌（手柄壳、摇杆帽、收纳包等），主要市场美国\n- Funlab（纷岚）：游戏外设品牌（手柄、配件），主要市场美国'

    wow_pk_spend = wow(pk_total['spend'], pk_last_total['spend'])
    wow_pk_roas = wow(pk_total['roas'], pk_last_total['roas'])
    wow_fl_spend = wow(fl_total['spend'], fl_last_total['spend'])
    wow_fl_roas = wow(fl_total['roas'], fl_last_total['roas'])

    # Enrich ad body text from Ad Library data (API often returns empty body)
    if ad_library_ads:
        lib_bodies = {la.get('body', '')[:40]: la.get('body', '') for la in ad_library_ads if la.get('body')}
        for ad in all_top_ads:
            if not ad.get('body'):
                # Try to find matching Ad Library body by any keyword
                for snippet, full_body in lib_bodies.items():
                    if snippet:  # assign first available if ad has no body
                        ad['body'] = full_body
                        del lib_bodies[snippet]
                        break

    top_detail = '\n'.join([
        f'{i+1}. [{ad["ad_name"]}] Account:{ad["account"]} | ROAS:{ad["roas"]} | '
        f'Spend:${ad["spend"]} | Purchases:{ad["purchases"]} | Type:{ad["creative_type"]} | '
        f'Campaign:{ad["campaign"]} | CPC:${ad["cpc"]} | CTR:{ad["ctr"]}%\n'
        f'   Ad Copy: {ad["body"][:200] if ad.get("body") else "(无文案数据)"}'
        for i, ad in enumerate(all_top_ads[:10])
    ])

    prompt = f'''===== 本周数据 =====

Powkong: Spend ${pk_total["spend"]:.2f} ({wow_pk_spend}) | ROAS {pk_total["roas"]} ({wow_pk_roas}) | Purchases {pk_total["purchases"]} | CPA ${pk_total["cpa"]}
上周: Spend ${pk_last_total["spend"]:.2f} | ROAS {pk_last_total["roas"]} | Purchases {pk_last_total["purchases"]}
Campaigns: {json.dumps(pk_campaigns, ensure_ascii=False)}

Funlab: Spend ${fl_total["spend"]:.2f} ({wow_fl_spend}) | ROAS {fl_total["roas"]} ({wow_fl_roas}) | Purchases {fl_total["purchases"]}
上周: Spend ${fl_last_total["spend"]:.2f} | ROAS {fl_last_total["roas"]}

===== Top 10 广告素材（含广告正文）=====
{top_detail}

===== 请严格基于system prompt中的知识库规则分析，按以下格式输出（每个板块用===分隔）=====

===整体表现总结===
2-3句概括双账户本周表现和关键变化

===Campaign排行===
按ROAS排名，标注盈利/亏损

===问题诊断===
2-3个具体问题，带数据。检查Campaign结构合规性、学习期阈值、受众策略

===Winning Ads素材分析===
分析表现最好的3-5条广告：每条分析停/病/药/信/买，标注漏斗层级（TOFU/MOFU/BOFU）

===本周假设验证===
哪种痛点角度ROAS最高？视频vs图片？哪种钩子CTR最高？

===下周内容制作指引===
1. 建议测试什么新痛点角度
2. 建议复制哪个Winning Ad结构
3. 建议制作几条什么类型素材

===下周行动项===
3-5条具体可执行待办'''

    try:
        ai_text = deepseek(system_prompt, prompt)
        log(f'  AI: {len(ai_text)} chars')
    except Exception as e:
        log(f'  DeepSeek failed: {e}')
        ai_text = f'分析失败: {e}'

    # Parse sections
    sections = {}
    current = ''
    for line in ai_text.split('\n'):
        match = re.match(r'={2,}\s*(.+?)\s*={2,}-?', line)
        if match:
            current = match.group(1)
            sections[current] = ''
        elif current:
            sections[current] += line + '\n'
    for k in sections:
        sections[k] = sections[k].strip()
    log(f'  Sections: {list(sections.keys())}')

    # ── Step 5: Build .docx ───────────────────────────────────
    log('Step 5: Building .docx with embedded images...')
    report_title = f'({date_str})META广告周报'
    docx_path = build_report_docx(
        report_title, pk_total, pk_last_total, fl_total, fl_last_total,
        pk_campaigns, fl_campaigns, all_top_ads, sections, ad_library_ads, now)
    log(f'  DOCX: {os.path.getsize(docx_path)} bytes')

    # ── Step 6: Import to Feishu wiki ─────────────────────────
    log('Step 6: Importing to Feishu wiki...')
    wiki_url, doc_token = import_docx_to_wiki(docx_path, report_title)
    try:
        os.remove(docx_path)
    except OSError:
        pass
    if not wiki_url:
        log('  Failed to create report!')
        return
    log(f'  Report: {wiki_url}')

    # ── Step 7: Notifications ─────────────────────────────────
    log('Step 7: Sending notifications...')
    card = {
        'config': {'wide_screen_mode': True},
        'header': {
            'title': {'tag': 'plain_text', 'content': report_title},
            'template': 'blue',
        },
        'elements': [
            {'tag': 'div', 'text': {'tag': 'lark_md',
                'content': (
                    f'**Powkong** | Spend: ${pk_total["spend"]:.0f} ({wow_pk_spend}) | ROAS: {pk_total["roas"]} | Purchases: {pk_total["purchases"]}\n'
                    f'**Funlab** | Spend: ${fl_total["spend"]:.0f} ({wow_fl_spend}) | ROAS: {fl_total["roas"]} | Purchases: {fl_total["purchases"]}'
                )}},
            {'tag': 'hr'},
            {'tag': 'action', 'actions': [{
                'tag': 'button',
                'text': {'tag': 'plain_text', 'content': '查看完整报告'},
                'url': wiki_url, 'type': 'primary',
            }]},
        ],
    }
    msg_body = json.dumps(card)
    recipients = [('Boss', BOSS_OPEN_ID)] + get_users_by_job_title('独立站运营专员')
    sent = set()
    for name, oid in recipients:
        if oid in sent:
            continue
        sent.add(oid)
        feishu('POST', '/im/v1/messages?receive_id_type=open_id',
               {'receive_id': oid, 'msg_type': 'interactive', 'content': msg_body})
        log(f'  Notified {name} ({oid})')

    # ── Write to system log ──
    try:
        duration = int((datetime.datetime.now() - now).total_seconds())
        feishu('POST', '/bitable/v1/apps/DLD5b93HLaWHVxs7NFjcisdgnuE/tables/tblk2ABjcQMnnxUh/records', {
            'fields': {
                '系统': 'S1广告周报',
                '执行时间': int(time.time() * 1000),
                '状态': '✅成功',
                '报告链接': {'link': wiki_url, 'text': report_title},
                '摘要': f'PK: ${pk_total["spend"]:.0f} ROAS={pk_total["roas"]} | FL: ${fl_total["spend"]:.0f} ROAS={fl_total["roas"]}',
                '耗时(秒)': duration,
            }
        })
    except Exception:
        pass

    log(f'\nS1 Weekly Report Complete!')
    log(f'  Report: {wiki_url}')

if __name__ == '__main__':
    main()
