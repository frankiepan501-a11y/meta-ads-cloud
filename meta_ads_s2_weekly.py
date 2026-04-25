"""S2 Competitor Monitoring - Production Weekly Script.
Triggered via FastAPI POST /run/s2 (cron orchestrated by n8n).
"""
import json, urllib.request, sys, time, re, datetime, uuid, os
try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

from fb_ad_library_scraper import scrape_ad_library

from cloud_config import (
    PROXY, IM_APP_ID, IM_APP_SECRET, BT_APP_ID, BT_APP_SECRET,
    DEEPSEEK_KEY, WIKI_SPACE_ID, BOSS_OPEN_ID, SKIP_ADLIB_IMAGES,
    META_ADS_CONSOLE_APP,
)

S2_APP_TOKEN = META_ADS_CONSOLE_APP
S2_TABLE_COMP = 'tblBniOwJYFYMDbQ'
S2_TABLE_ADS = 'tblyUsxc3NmHGAOZ'
S2_TABLE_WEEKLY = 'tbl8DPF9Z1jqdSpF'
WIKI_PARENT_NODE = os.environ.get('WIKI_PARENT_NODE_S2', 'Z41Awia1yiSTKAkUH8ScmBSHn6b')
LOG_FILE = os.environ.get('S2_LOG_FILE', '/tmp/s2_log.txt')

# Doc generation: build .docx with python-docx (images embedded in file) → upload → import → wiki.
# This avoids Feishu's import server needing to fetch external image URLs.

def log(msg):
    ts = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    line = f'[{ts}] {msg}'
    print(line)
    with open(LOG_FILE, 'a', encoding='utf-8') as f:
        f.write(line + '\n')

# ============================================================
# API helpers (with token caching)
# ============================================================
_token_cache = {}

def get_token(app_id, secret):
    key = app_id
    if key in _token_cache and time.time() - _token_cache[key][1] < 600:
        return _token_cache[key][0]
    body = json.dumps({'app_id': app_id, 'app_secret': secret}).encode()
    req = urllib.request.Request(
        'https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal',
        data=body, headers={'Content-Type': 'application/json'})
    token = json.loads(urllib.request.build_opener(PROXY).open(req, timeout=15).read())['tenant_access_token']
    _token_cache[key] = (token, time.time())
    return token

def get_users_by_job_title(title):
    r = feishu('GET', '/contact/v3/departments?parent_department_id=0&page_size=50'
               '&department_id_type=open_department_id&user_id_type=open_id&fetch_child=true',
               app='im')
    dept_ids = ['0'] + [d['open_department_id'] for d in r.get('data', {}).get('items', [])]
    seen, matched = set(), []
    for did in dept_ids:
        page = ''
        while True:
            q = (f'/contact/v3/users/find_by_department?department_id={did}'
                 f'&page_size=50&user_id_type=open_id&department_id_type=open_department_id')
            if page:
                q += f'&page_token={page}'
            r = feishu('GET', q, app='im')
            for u in r.get('data', {}).get('items', []):
                oid = u.get('open_id')
                if not oid or oid in seen:
                    continue
                seen.add(oid)
                if u.get('job_title') == title:
                    matched.append((u.get('name'), oid))
            page = r.get('data', {}).get('page_token', '')
            if not page:
                break
    return matched

def feishu(method, path, body=None, app='bt'):
    url = f'https://open.feishu.cn/open-apis{path}'
    data = json.dumps(body).encode('utf-8') if body else None
    last_err = None
    for attempt in range(3):
        token = get_token(BT_APP_ID, BT_APP_SECRET) if app == 'bt' else get_token(IM_APP_ID, IM_APP_SECRET)
        req = urllib.request.Request(url, data=data,
            headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {token}'}, method=method)
        try:
            return json.loads(urllib.request.build_opener(PROXY).open(req, timeout=30).read())
        except urllib.error.HTTPError as e:
            err = e.read().decode('utf-8', 'replace')[:500]
            log(f'  API Error {e.code}: {err[:200]}')
            return {'code': e.code, 'body': err}
        except (urllib.error.URLError, OSError) as e:
            last_err = e
            log(f'  Network error (attempt {attempt+1}/3): {e}')
            time.sleep(2 * (attempt + 1))
    return {'code': -1, 'body': f'Network error after 3 retries: {last_err}'}

def deepseek(system_prompt, user_prompt):
    body = json.dumps({
        'model': 'deepseek-chat', 'max_tokens': 6000,
        'messages': [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt},
        ]
    }).encode('utf-8')
    req = urllib.request.Request('https://api.deepseek.com/chat/completions',
        data=body, headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {DEEPSEEK_KEY}'})
    resp = json.loads(urllib.request.build_opener(PROXY).open(req, timeout=120).read())
    return resp['choices'][0]['message']['content']


# ============================================================
# Week-over-week comparison
# ============================================================
def fetch_existing_ads(comp_name):
    """Return {ad_id: record_id} for ads currently marked as 仍在投放."""
    resp = feishu('POST',
        f'/bitable/v1/apps/{S2_APP_TOKEN}/tables/{S2_TABLE_ADS}/records/search?page_size=100',
        {'filter': {'conjunction': 'and', 'conditions': [
            {'field_name': '竞品名称', 'operator': 'is', 'value': [comp_name]},
            {'field_name': '仍在投放', 'operator': 'is', 'value': ['true']},
        ]}, 'field_names': ['广告ID', '竞品名称', '仍在投放']})
    result = {}
    for item in resp.get('data', {}).get('items', []):
        v = item['fields'].get('广告ID', '')
        ad_id = v if isinstance(v, str) else (v[0].get('text', '') if isinstance(v, list) and v and isinstance(v[0], dict) else '')
        if ad_id:
            result[ad_id] = item['record_id']
    return result

def mark_ads_stopped(record_ids):
    """Set 仍在投放=false for given record IDs."""
    for rid in record_ids:
        feishu('PUT', f'/bitable/v1/apps/{S2_APP_TOKEN}/tables/{S2_TABLE_ADS}/records/{rid}',
               {'fields': {'仍在投放': False}})

# ============================================================
# Image download helper
# ============================================================
def download_image(url):
    """Download image from URL and compress for docx embedding. Returns bytes or None."""
    try:
        req = urllib.request.Request(url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'})
        data = urllib.request.build_opener(PROXY).open(req, timeout=15).read()
        if len(data) < 500:
            return None
    except Exception:
        return None
    try:
        from PIL import Image
        import io as _io
        im = Image.open(_io.BytesIO(data))
        if im.width > 1000:
            ratio = 1000 / im.width
            im = im.resize((1000, int(im.height * ratio)))
        out = _io.BytesIO()
        im.convert('RGB').save(out, format='JPEG', quality=78, optimize=True)
        return out.getvalue()
    except Exception:
        return data

# ============================================================
# Build .docx report with embedded images → import to Feishu
# ============================================================
def build_docx_report(all_reports, week_str, now):
    """Build a .docx file with embedded ad images. Returns file path."""
    from docx import Document
    from docx.shared import Inches
    import io, tempfile

    doc = Document()

    # Header
    date_str = now.strftime('%Y-%m-%d')
    doc.add_heading(f'({date_str})竞品META广告分析', level=1)
    doc.add_paragraph(f'生成时间: {now.strftime("%Y-%m-%d %H:%M")}')

    # Overview
    doc.add_heading('本周总览', level=2)
    for rpt in all_reports:
        doc.add_paragraph(
            f'{rpt["name"]} — 活跃:{rpt["ads_count"]} | 新增:{rpt["new_count"]}'
            f' | 停投:{rpt["stopped_count"]} | 威胁:{rpt["threat"]}',
            style='List Bullet')

    # Per-competitor
    for rpt in all_reports:
        doc.add_heading(f'{rpt["name"]}（威胁：{rpt["threat"]}）', level=2)
        p = doc.add_paragraph()
        p.add_run(f'活跃广告: {rpt["ads_count"]}').bold = True
        p.add_run(f' | 新增: {rpt["new_count"]} | 停投: {rpt["stopped_count"]}'
                  f' | 持续: {rpt["continuing_count"]}')

        doc.add_heading('广告拆解', level=3)

        for i, ad in enumerate(rpt['ads'][:10]):
            analysis = rpt['per_ad'][i] if i < len(rpt['per_ad']) else {}
            ad_id = ad.get('id', '')
            is_new = ad_id in rpt['new_ids']
            tag = ' [新]' if is_new else ''

            doc.add_heading(
                f'▎广告 #{i+1}{tag} — {ad.get("pageName", "")}', level=4)

            # Embed ad image (skipped on cloud v1; use Ad Library link instead)
            images = ad.get('images', [])
            if SKIP_ADLIB_IMAGES:
                if ad_id:
                    p = doc.add_paragraph()
                    p.add_run('Ad Library: ').bold = True
                    p.add_run(f'https://www.facebook.com/ads/library/?id={ad_id}')
            elif images and i < 8:
                img_data = download_image(images[0])
                if img_data:
                    try:
                        doc.add_picture(io.BytesIO(img_data), width=Inches(5))
                    except Exception as e:
                        log(f'    Skip image #{i+1}: {e}')
                        doc.add_paragraph(f'[图片加载失败: {images[0][:60]}...]')

            body_text = ad.get('body', '')[:300].replace('\n', ' ')
            if body_text:
                p = doc.add_paragraph()
                p.add_run('Ad Copy: ').bold = True
                p.add_run(body_text)

            cta = ad.get('cta', '')
            if cta:
                p = doc.add_paragraph()
                p.add_run('CTA: ').bold = True
                p.add_run(cta)

            # AI analysis
            for label, field in [('停(Hook)', 'stop'), ('病(Pain)', 'pain'),
                                 ('药(Solution)', 'medicine'), ('信(Trust)', 'trust'),
                                 ('买(CTA)', 'buy')]:
                val = analysis.get(field, '')
                if val:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{label}: ').bold = True
                    p.add_run(val)

            rating = analysis.get('rating', '')
            if rating:
                doc.add_paragraph(f'评分: {rating}')

            borrow = analysis.get('borrow', '')
            if borrow:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run('💡 借鉴: ').bold = True
                p.add_run(borrow)

            improve = analysis.get('improve', '')
            if improve:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run('🔧 改进: ').bold = True
                p.add_run(improve)

            if ad_id:
                doc.add_paragraph(
                    f'→ 查看原广告: https://www.facebook.com/ads/library/?id={ad_id}')

        # Overall sections
        for key, title in [('整体策略', '策略洞察'), ('周变化点评', '本周变化'),
                           ('威胁评估', '威胁评估'), ('行动建议', '行动建议')]:
            content = (rpt['overall'].get(key, '') or '').strip()
            if not content:
                continue
            doc.add_heading(title, level=3)
            for line in content.split('\n')[:20]:
                t = line.strip()
                if not t:
                    continue
                clean = re.sub(r'^[\d\.\)、\-*•]+\s*', '', t)
                if t.startswith(('- ', '* ', '• ')) or re.match(r'^\d+[\.\)、]', t):
                    doc.add_paragraph(clean, style='List Bullet')
                else:
                    doc.add_paragraph(t)

    # Action items
    doc.add_heading('汇总行动项', level=2)
    for rpt in all_reports:
        actions = (rpt['overall'].get('行动建议', '') or '').strip()
        if not actions:
            continue
        doc.add_heading(rpt['name'], level=3)
        for line in actions.split('\n')[:8]:
            t = line.strip()
            if not t:
                continue
            clean = re.sub(r'^[\d\.\)、\-*•]+\s*', '', t)
            if clean:
                doc.add_paragraph(f'☐ {clean}', style='List Bullet')

    path = os.path.join(tempfile.gettempdir(), f's2_report_{week_str}.docx')
    doc.save(path)
    return path

def _upload_and_import_docx(docx_bytes, title):
    """One attempt: upload docx + import + poll. Returns doc_token or None."""
    token = get_token(IM_APP_ID, IM_APP_SECRET)
    boundary = uuid.uuid4().hex
    safe_name = f'{title}.docx'
    extra_json = json.dumps({'obj_type': 'docx', 'file_extension': 'docx'})
    parts = []
    for name, value in [('file_name', safe_name), ('parent_type', 'ccm_import_open'),
                        ('size', str(len(docx_bytes))), ('extra', extra_json)]:
        parts.append(f'--{boundary}\r\nContent-Disposition: form-data; name="{name}"\r\n\r\n{value}\r\n'.encode())
    parts.append(f'--{boundary}\r\nContent-Disposition: form-data; name="file"; filename="{safe_name}"\r\n'
                 f'Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n'.encode())
    parts.append(docx_bytes)
    parts.append(f'\r\n--{boundary}--\r\n'.encode())

    req = urllib.request.Request('https://open.feishu.cn/open-apis/drive/v1/medias/upload_all',
        data=b''.join(parts),
        headers={'Content-Type': f'multipart/form-data; boundary={boundary}', 'Authorization': f'Bearer {token}'})
    try:
        resp = json.loads(urllib.request.build_opener(PROXY).open(req, timeout=60).read())
        file_token = resp['data']['file_token']
    except Exception as e:
        log(f'  DOCX upload failed: {e}')
        return None

    r = feishu('POST', '/drive/v1/import_tasks', {
        'file_extension': 'docx', 'file_token': file_token,
        'type': 'docx', 'point': {'mount_type': 1, 'mount_key': ''},
    }, app='im')
    if r.get('code') != 0:
        log(f'  Import task failed: {r}')
        return None
    ticket = r['data']['ticket']

    for _ in range(60):
        time.sleep(2)
        r2 = feishu('GET', f'/drive/v1/import_tasks/{ticket}', app='im')
        status = r2.get('data', {}).get('result', {}).get('job_status', -1)
        if status == 0:
            return r2['data']['result'].get('token', '')
        elif status in (1, 2):
            result = r2.get('data', {}).get('result', {})
            log(f'  Import failed: status={status} msg={result.get("job_error_msg","")} '
                f'extra={json.dumps(result, ensure_ascii=False)[:300]}')
            return None
    log('  Import timeout')
    return None


def import_docx_to_wiki(docx_path, title):
    """Upload .docx, import as Feishu doc, move to wiki. Returns (wiki_url, doc_token)."""
    with open(docx_path, 'rb') as f:
        docx_bytes = f.read()
    log(f'  DOCX size: {len(docx_bytes)} bytes')

    doc_token = None
    for attempt in range(2):
        doc_token = _upload_and_import_docx(docx_bytes, title)
        if doc_token:
            break
        log(f'  Import attempt {attempt+1}/2 failed, retrying...')
        time.sleep(5)
    if not doc_token:
        return None, None

    # Permissions (before move, else app loses move permission after transfer)
    feishu('POST', f'/drive/v1/permissions/{doc_token}/members?type=docx&need_notification=false',
           {'member_type': 'openid', 'member_id': BOSS_OPEN_ID, 'perm': 'full_access'}, app='im')

    # Move to wiki
    r3 = feishu('POST', f'/wiki/v2/spaces/{WIKI_SPACE_ID}/nodes/move_docs_to_wiki', {
        'parent_wiki_token': WIKI_PARENT_NODE,
        'obj_type': 'docx', 'obj_token': doc_token,
    }, app='im')
    if r3.get('code') != 0:
        log(f'  Move to wiki failed: {r3}')
        return None, None

    node_token = r3.get('data', {}).get('wiki_token')
    if not node_token:
        for _ in range(10):
            time.sleep(2)
            rn = feishu('GET', f'/wiki/v2/spaces/get_node?obj_type=docx&token={doc_token}', app='im')
            node_token = rn.get('data', {}).get('node', {}).get('node_token')
            if node_token:
                break
        if not node_token:
            log(f'  Wiki node_token not resolved: {r3}')
            return None, None

    wiki_url = f'https://u1wpma3xuhr.feishu.cn/wiki/{node_token}'
    return wiki_url, doc_token

# ============================================================
# Parse per-ad analysis from DeepSeek output
# ============================================================
def parse_per_ad_analysis(text):
    """Parse structured per-ad analysis + overall sections.
    Returns (list_of_ad_dicts, dict_of_overall_sections).
    """
    ads = []
    cur_ad = None
    overall = {}
    cur_section = ''
    in_overall = False

    for line in text.split('\n'):
        ls = line.strip()

        # Overall section markers: === ... === or ## ...
        m = re.match(r'={2,}\s*(.+?)\s*={2,}-?', ls)
        if not m:
            m = re.match(r'^##\s+(.+)', ls)
        if m:
            name = m.group(1).strip()
            if cur_ad:
                ads.append(cur_ad); cur_ad = None
            cur_section = name
            overall[cur_section] = ''
            in_overall = True
            continue

        # Per-ad marker: 【广告#N】
        m2 = re.match(r'【广告[#＃]?\s*(\d+)】', ls)
        if m2:
            if cur_ad:
                ads.append(cur_ad)
            cur_ad = {}
            in_overall = False
            continue

        if cur_ad is not None and not in_overall:
            # Parse fields within an ad analysis
            for key, field in [('停', 'stop'), ('病', 'pain'), ('药', 'medicine'),
                               ('信', 'trust'), ('买', 'buy'), ('评分', 'rating'),
                               ('借鉴', 'borrow'), ('改进', 'improve')]:
                pat = rf'^[·•\-*\s]*{key}[\(（][^)）]*[\)）]?\s*[:：]\s*(.+)'
                m = re.match(pat, ls)
                if not m:
                    m = re.match(rf'^[·•\-*\s]*{key}\s*[:：]\s*(.+)', ls)
                if m:
                    cur_ad[field] = m.group(1).strip()
                    break
        elif in_overall and cur_section:
            overall[cur_section] += line + '\n'

    if cur_ad:
        ads.append(cur_ad)
    return ads, overall

# ============================================================
# Main
# ============================================================
def main():
    log('=' * 60)
    log('S2 Competitor Monitoring v2 - Weekly Run')

    now = datetime.datetime.now()
    fake_today = os.environ.get('META_FAKE_TODAY')
    if fake_today:
        now = datetime.datetime.strptime(fake_today, '%Y-%m-%d')
        log(f'  [OVERRIDE] today = {fake_today}')
    week_str = f'{now.year}-W{now.isocalendar()[1]:02d}'

    kb_path = os.environ.get('META_ADS_KB_PATH', os.path.join(os.path.dirname(os.path.abspath(__file__)), 'meta_ads_knowledge.md'))
    with open(kb_path, 'r', encoding='utf-8') as f:
        kb = f.read()

    # ── Step 1: Read competitors ──────────────────────────────
    log('Step 1: Reading competitors from Feishu...')
    resp = feishu('POST',
        f'/bitable/v1/apps/{S2_APP_TOKEN}/tables/{S2_TABLE_COMP}/records/search?page_size=20',
        {'filter': {'conjunction': 'and', 'conditions': [
            {'field_name': '监控状态', 'operator': 'is', 'value': ['监控中']}
        ]}})
    competitors = []
    for item in resp.get('data', {}).get('items', []):
        fd = item['fields']
        def gt(v):
            if not v: return ''
            if isinstance(v, str): return v
            if isinstance(v, list) and v and isinstance(v[0], dict): return v[0].get('text', '')
            if isinstance(v, dict) and 'link' in v: return v['link']
            return str(v)
        name = gt(fd.get('竞品名称', ''))
        if name:
            # Read exact Facebook page names for strict matching
            page_names_raw = gt(fd.get('Facebook页面名称', ''))
            page_names = [p.strip() for p in page_names_raw.split(',') if p.strip()]
            competitors.append({
                'name': name,
                'category': gt(fd.get('品类', '')),
                'record_id': item['record_id'],
                'page_names': page_names,  # exact FB page names
            })
    log(f'  Found {len(competitors)} competitors: {[c["name"] for c in competitors]}')

    all_reports = []

    # ── Step 2-4: Process each competitor ─────────────────────
    for comp in competitors:
        log(f'\n--- Processing [{comp["name"]}] ---')

        # 2a. WoW: fetch existing active ads
        log('  Fetching existing ads for WoW comparison...')
        existing_ads = fetch_existing_ads(comp['name'])
        log(f'  Existing active: {len(existing_ads)}')

        # 2b. Scrape Ad Library
        log('  Scraping Ad Library...')
        try:
            all_ads = scrape_ad_library(comp['name'], 'ALL', 20)
            log(f'  Total scraped: {len(all_ads)}')

            # Strict page name matching (if configured)
            if comp['page_names']:
                pn_lower = [p.lower() for p in comp['page_names']]
                official = [a for a in all_ads
                            if a.get('pageName', '').lower() in pn_lower]
                log(f'  Strict page match ({comp["page_names"]}): {len(official)}')
            else:
                # Fallback: fuzzy match on brand name in page name
                official = [a for a in all_ads
                            if comp['name'].lower() in a.get('pageName', '').lower()]
                log(f'  Fuzzy page match: {len(official)}')
        except Exception as e:
            log(f'  Scrape failed: {e}')
            all_ads = []; official = []

        # Deduplicate ads with identical body text
        if official:
            seen_bodies = {}
            deduped = []
            for a in official:
                body_key = a.get('body', '')[:80].strip()
                if not body_key or body_key not in seen_bodies:
                    deduped.append(a)
                    if body_key:
                        seen_bodies[body_key] = 1
                else:
                    seen_bodies[body_key] += 1
            dupes = len(official) - len(deduped)
            if dupes:
                log(f'  Deduped: {len(official)} → {len(deduped)} (removed {dupes} duplicates)')
            official = deduped

        if not official:
            log('  No matching ads found, skipping')

        if not official:
            log('  No ads at all, skipping')
            feishu('PUT',
                f'/bitable/v1/apps/{S2_APP_TOKEN}/tables/{S2_TABLE_COMP}/records/{comp["record_id"]}',
                {'fields': {'上次扫描': int(time.time() * 1000), '活跃广告数': 0}})
            continue

        # 2c. WoW diff
        # Fix: scraper returns 'id' not 'ad_archive_id'
        current_ids = {a.get('id', '') for a in official if a.get('id')}
        new_ids = current_ids - set(existing_ads.keys())
        stopped_ids = set(existing_ads.keys()) - current_ids
        continuing_ids = current_ids & set(existing_ads.keys())
        log(f'  WoW: new={len(new_ids)} stopped={len(stopped_ids)} cont={len(continuing_ids)}')

        # Mark stopped ads in bitable
        if stopped_ids:
            stopped_rids = [existing_ads[aid] for aid in stopped_ids if aid in existing_ads]
            mark_ads_stopped(stopped_rids)
            log(f'  Marked {len(stopped_rids)} ads as stopped')

        # 2d. DeepSeek per-ad analysis
        log('  Calling DeepSeek (per-ad analysis)...')
        ads_text = '\n'.join([
            f'【广告#{i+1}】\n'
            f'广告主: {a.get("pageName","")}\n'
            f'Ad Copy: {a["body"][:300]}\n'
            f'CTA: {a.get("cta","")}\n'
            f'投放开始: {a.get("startDate","")}\n'
            f'素材类型: {a.get("mediaType","image")}\n'
            f'{"[新广告]" if a.get("id","") in new_ids else "[持续投放]"}'
            for i, a in enumerate(official[:15])
        ])

        wow_hint = ''
        if new_ids or stopped_ids:
            wow_hint = (f'\n\n本周变化：新增{len(new_ids)}条，停投{len(stopped_ids)}条，'
                        f'持续投放{len(continuing_ids)}条。请在整体策略中点评变化含义。')

        sys_prompt = (kb + f'\n\n品牌信息：分析竞品【{comp["name"]}】'
                      f'（品类：{comp["category"]}）的广告策略。')
        usr_prompt = f'''以下是竞品【{comp["name"]}】当前 {len(official)} 条 Facebook 活跃广告：

{ads_text}{wow_hint}

请严格按以下格式输出：

**第一部分：逐条广告拆解**（每条广告用【广告#N】开头）

【广告#1】
停(Hook): 前3秒钩子类型和话术
病(Pain): 痛点/情绪分析
药(Solution): 卖点展示方式
信(Trust): 社会证明（无则写"无"）
买(CTA): 促单话术和权益
评分: ⭐⭐⭐ (1-5颗星)
借鉴: 值得学习的点
改进: 如果我们做类似内容的改进方向

【广告#2】
... （同上格式）

**第二部分：整体分析**（每个板块用 === 包裹标题）

===整体策略===
投放模式、内容方向、受众策略总结

===周变化点评===
新增/停投广告的含义，策略调整方向

===威胁评估===
对我们的威胁等级（低/中/高）和具体原因

===行动建议===
我们可以采取的 3-5 条具体行动'''

        try:
            ai_text = deepseek(sys_prompt, usr_prompt)
            log(f'  AI: {len(ai_text)} chars')
        except Exception as e:
            log(f'  DeepSeek failed: {e}')
            ai_text = f'分析失败: {e}'

        per_ad, overall = parse_per_ad_analysis(ai_text)
        log(f'  Parsed: {len(per_ad)} ads, {len(overall)} sections')

        # 2e. Write ads to bitable (with per-ad analysis fields)
        log('  Writing ads to bitable...')
        for idx, ad in enumerate(official[:10]):
            ad_id = ad.get('id', '')
            # Skip duplicates that already exist
            if ad_id and ad_id in existing_ads:
                continue

            analysis = per_ad[idx] if idx < len(per_ad) else {}
            fields = {
                '竞品名称': comp['name'],
                '广告ID': ad_id[:50],
                'Ad Copy': ad['body'][:2000],
                'Headline': (ad['body'].split('\n')[0] if ad['body'] else '')[:200],
                '素材类型': '视频' if ad.get('hasVideo') else '图片',
                '首次发现日期': int(time.time() * 1000),
                '仍在投放': True,
            }
            if ad.get('cta'):
                fields['买(CTA)'] = ad['cta']
            # Write per-ad AI analysis fields
            if analysis.get('stop'):
                fields['停(Hook)'] = analysis['stop'][:500]
            if analysis.get('pain'):
                fields['病(Pain)'] = analysis['pain'][:500]
            if analysis.get('medicine'):
                fields['药(Solution)'] = analysis['medicine'][:500]
            if analysis.get('trust'):
                fields['信(Trust)'] = analysis['trust'][:500]
            if analysis.get('improve'):
                fields['改进建议'] = analysis['improve'][:500]
            # Build AI分析 summary
            ai_summary = '\n'.join(
                f'{k}: {analysis.get(v, "")}'
                for k, v in [('评分', 'rating'), ('借鉴', 'borrow'), ('改进', 'improve')]
                if analysis.get(v)
            )
            if ai_summary:
                fields['AI 分析'] = ai_summary[:2000]

            # Ad Library permalink
            if ad_id:
                fields['广告帖链接'] = {
                    'link': f'https://www.facebook.com/ads/library/?id={ad_id}',
                    'text': f'Ad Library #{ad_id}',
                }

            feishu('POST',
                f'/bitable/v1/apps/{S2_APP_TOKEN}/tables/{S2_TABLE_ADS}/records',
                {'fields': fields})

        # 2f. Write weekly summary to bitable
        threat = '低'
        t_text = overall.get('威胁评估', '')
        if '高' in t_text:
            threat = '高'
        elif '中' in t_text:
            threat = '中'

        weekly = {
            '周次': week_str,
            '竞品名称': comp['name'],
            '报告日期': int(time.time() * 1000),
            '本周新增广告数': len(new_ids) if new_ids else len(official),
            '本周停投广告数': len(stopped_ids),
            '策略洞察': (overall.get('整体策略', '') or ai_text[:2000]).strip()[:2000],
            '值得借鉴的角度': (overall.get('行动建议', '') or '').strip()[:2000],
        }
        r = feishu('POST',
            f'/bitable/v1/apps/{S2_APP_TOKEN}/tables/{S2_TABLE_WEEKLY}/records',
            {'fields': weekly})
        if r.get('code') == 0:
            rid = r['data']['record']['record_id']
            feishu('PUT',
                f'/bitable/v1/apps/{S2_APP_TOKEN}/tables/{S2_TABLE_WEEKLY}/records/{rid}',
                {'fields': {'威胁评估': threat}})

        # Update competitor record
        feishu('PUT',
            f'/bitable/v1/apps/{S2_APP_TOKEN}/tables/{S2_TABLE_COMP}/records/{comp["record_id"]}',
            {'fields': {'上次扫描': int(time.time() * 1000), '活跃广告数': len(official)}})

        all_reports.append({
            'name': comp['name'],
            'ads': official,
            'ads_count': len(official),
            'new_ids': new_ids,
            'stopped_ids': stopped_ids,
            'new_count': len(new_ids),
            'stopped_count': len(stopped_ids),
            'continuing_count': len(continuing_ids),
            'threat': threat,
            'per_ad': per_ad,
            'overall': overall,
            'ai_text': ai_text,
        })
        log(f'  Done: {len(official)} ads (new:{len(new_ids)} stop:{len(stopped_ids)}) threat={threat}')

    # ── Step 5: Build .docx with images and import ──────────
    if not all_reports:
        log('No reports to generate.')
        return

    log('\nStep 5: Building .docx report with embedded images...')
    docx_path = build_docx_report(all_reports, week_str, now)
    log(f'  DOCX: {os.path.getsize(docx_path)} bytes')

    date_str = now.strftime('%Y-%m-%d')
    wiki_url, doc_token = import_docx_to_wiki(docx_path, f'({date_str})竞品META广告分析')
    try:
        os.remove(docx_path)
    except OSError:
        pass
    if not wiki_url:
        log('  Failed to create report doc!')
        return
    log(f'  Report: {wiki_url}')

    # ── Notifications ─────────────────────────────────────────
    summary_lines = [
        f'**{r["name"]}** | 活跃:{r["ads_count"]} 新增:{r["new_count"]}'
        f' 停投:{r["stopped_count"]} | 威胁:{r["threat"]}'
        for r in all_reports
    ]
    card = {
        'config': {'wide_screen_mode': True},
        'header': {
            'title': {'tag': 'plain_text', 'content': f'({date_str})竞品META广告分析'},
            'template': 'purple',
        },
        'elements': [
            {'tag': 'div', 'text': {'tag': 'lark_md', 'content': '\n'.join(summary_lines)}},
            {'tag': 'hr'},
            {'tag': 'action', 'actions': [{
                'tag': 'button',
                'text': {'tag': 'plain_text', 'content': '查看完整报告'},
                'url': wiki_url, 'type': 'primary',
            }]},
        ],
    }
    msg_body = json.dumps(card)
    recipients = [('Boss', BOSS_OPEN_ID)] + get_users_by_job_title('独立站运营专员')
    sent = set()
    for name, oid in recipients:
        if oid in sent:
            continue
        sent.add(oid)
        feishu('POST', '/im/v1/messages?receive_id_type=open_id',
               {'receive_id': oid, 'msg_type': 'interactive', 'content': msg_body},
               app='im')
        log(f'  Notified {name} ({oid})')

    # ── Write to system log ──
    try:
        duration = int((datetime.datetime.now() - now).total_seconds())
        comp_summary = ' | '.join(f'{r["name"]}:{r["ads_count"]}ads' for r in all_reports)
        feishu('POST', '/bitable/v1/apps/DLD5b93HLaWHVxs7NFjcisdgnuE/tables/tblk2ABjcQMnnxUh/records', {
            'fields': {
                '系统': 'S2竞品监控',
                '执行时间': int(time.time() * 1000),
                '状态': '✅成功',
                '报告链接': {'link': wiki_url, 'text': f'竞品分析-{week_str}'},
                '摘要': comp_summary,
                '耗时(秒)': duration,
            }
        }, app='im')
    except Exception:
        pass

    log(f'\nS2 v2 Weekly Run Complete! {len(all_reports)} competitors processed.')
    log(f'  Report: {wiki_url}')

if __name__ == '__main__':
    main()
